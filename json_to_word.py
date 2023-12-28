# so this program will take a json file and convert it to a word file
# the json file will be a list of dictionaries
# each dictionary will have a key called "question" which is the question
# then they will contain a list of answers and points the answers are worth the points are in the same order as the answers
# now we want to convert this to a word file
# where the question is the question
# then the answers and points in 2 column the answers on the left and the points on the right
# then a line break

# import the json library
import json
import os
import sys
import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

file_name = "questionBank.json"

# open the json file

with open(file_name) as json_file:
    data = json.load(json_file)
    print(data)

    # now we want to create a word document
    document = Document()

    # now we want to add a title
    document.add_heading('Question Bank', 0)

    # now we parse the json file
    for question in data:
        # now we want to add the question
        document.add_heading(question["question"], level=1)
        # now we want to add the answers
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Answer'
        hdr_cells[1].text = 'Points'
        for answer, points in zip(question["answers"], question["points"]):
            row_cells = table.add_row().cells
            row_cells[0].text = answer
            row_cells[1].text = str(points)

    # now we want to save the document
    document.save('questionBank.docx')